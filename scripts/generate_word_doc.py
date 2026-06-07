"""
Generates DevOps Case Study Word document (.docx)
Output: DevOps-CaseStudy-Submission.docx
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SS   = os.path.join(BASE, 'screenshots')
OUT  = os.path.join(BASE, 'DevOps-CaseStudy-Submission.docx')

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── helpers ──────────────────────────────────────────────────────────────────
def add_toc(doc):
    """Insert a real Word TOC field — clickable hyperlinks, auto page numbers.
    When the doc is opened in Word, press Ctrl+A then F9 to refresh page numbers,
    or right-click the TOC area → Update Field."""
    # Heading
    th = doc.add_heading('Table of Contents', level=1)
    th.runs[0].font.color.rgb = RGBColor(0x32, 0x6C, 0xE5)

    # Instruction paragraph
    ip = doc.add_paragraph()
    ir = ip.add_run(
        'When opened in Word: press Ctrl+A then F9 to update page numbers and links.'
    )
    ir.italic = True
    ir.font.size = Pt(9)
    ir.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

    # TOC field paragraph
    p = doc.add_paragraph()

    # ── begin field ──
    run = p.add_run()
    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(qn('w:fldCharType'), 'begin')
    fc_begin.set(qn('w:dirty'), 'true')   # forces update on open
    run._r.append(fc_begin)

    # ── instruction text ──
    #   \o "1-3"  include Heading 1-3
    #   \h         make entries hyperlinks
    #   \z         hide tab/page in web view
    #   \u         use outline level
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)

    # ── separate ──
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fc_sep)

    # ── placeholder (replaced by Word on update) ──
    run2 = p.add_run()
    ph = OxmlElement('w:t')
    ph.text = '[Right-click here → Update Field to build the TOC]'
    run2._r.append(ph)
    run2.font.size = Pt(10)
    run2.italic = True
    run2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── end field ──
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    run2._r.append(fc_end)

def h1(text, color=(0x32,0x6C,0xE5)):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(*color)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def para(text='', bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold; run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level*0.5)
    return p

def numbered(text):
    return doc.add_paragraph(text, style='List Number')

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'),   'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'),  'F3F4F6')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1F,0x29,0x37)
    return p

def img(path, width=Inches(5.8), caption_text=''):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption_text:
            cp = doc.add_paragraph(caption_text)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].font.size = Pt(9)
            cp.runs[0].italic = True
            cp.runs[0].font.color.rgb = RGBColor(0x6C,0x75,0x7D)
    else:
        para(f'[Screenshot not found: {os.path.basename(path)}]',
             italic=True, color=(180,180,180))

def blue_header_table(headers, rows, col_color='326CE5'):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = t.rows[0].cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'), col_color)
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            t.rows[ri+1].cells[ci].text = str(val)
            t.rows[ri+1].cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
    return t

def page_break():
    doc.add_page_break()

def divider():
    doc.add_paragraph('─'*80).runs[0].font.color.rgb = RGBColor(0xDE,0xE2,0xE6)

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
for _ in range(4): doc.add_paragraph()

for txt, sz, col in [
    ('DevOps Case Study', 28, (0x32,0x6C,0xE5)),
    ('CSV File Processor — End-to-End DevOps Implementation', 16, (0x1A,0x1A,0x2E)),
    ('', 12, (0x6C,0x75,0x7D)),
    ('Kubernetes (kops) · Helm · Ansible · Terraform · Minikube', 12, (0x6C,0x75,0x7D)),
    ('Docker · MinIO · AWS S3 + Glacier · Nginx · Flask · Python', 11, (0x6C,0x75,0x7D)),
    ('', 10, (0x6C,0x75,0x7D)),
    ('Candidate:  Deepak Inugala', 12, (0x6C,0x75,0x7D)),
    ('Image:      deepak415/csv-processor:latest  (DockerHub)', 11, (0x6C,0x75,0x7D)),
    ('Repository: github.com/AIINDEVOPS/silk', 11, (0x6C,0x75,0x7D)),
    ('Date:       June 2026', 11, (0x6C,0x75,0x7D)),
]:
    lp = doc.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run(txt)
    lr.font.size = Pt(sz)
    if sz == 28: lr.bold = True
    if sz == 16: lr.bold = True
    lr.font.color.rgb = RGBColor(*col)

for _ in range(5): doc.add_paragraph()
page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS  (real Word TOC field — clickable hyperlinks)
# ════════════════════════════════════════════════════════════════════════════
add_toc(doc)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — REQUIREMENTS
# ════════════════════════════════════════════════════════════════════════════
h1('1. Assignment Requirements — Overview')
para(
    'The assignment has two parts: (1) a Kubernetes cluster configuration using kops '
    '— a running cluster is not expected, only the config files; and (2) a working '
    'application deployed on Minikube locally, covering all infrastructure, '
    'application, and storage requirements.'
)
doc.add_paragraph()

blue_header_table(
    ['Requirement (from assignment)', 'Status', 'Location'],
    [
        ('kops cluster config: multiple IGs, spot+on-demand', 'Done', 'k8s-kops/instancegroups.yaml'),
        ('Cluster autoscaler for all instance groups', 'Done', 'k8s-kops/cluster-autoscaler.yaml'),
        ('Nginx + web app in same pod, shared storage (not NFS)', 'Done', 'app/nginx.conf, helm/csv-app/templates/deployment.yaml'),
        ('Expose application with Service object', 'Done', 'helm/csv-app/templates/service.yaml'),
        ('Implement autoscaling for deployment', 'Done', 'helm/csv-app/templates/hpa.yaml'),
        ('Ansible for application configs', 'Done', 'ansible/site.yaml'),
        ('Helm for Kubernetes objects reuse across environments', 'Done', 'helm/csv-app/ + helm/environments/'),
        ('Python web app to parse and process CSV files', 'Done', 'app/app.py'),
        ('Upload CSV interface + show previously processed files', 'Done', 'app/templates/index.html'),
        ('Once processed, upload to S3 storage', 'Done', 'app/app.py (boto3 → MinIO/S3)'),
        ('Implement S3 Glacier transition on S3 config', 'Done', 'terraform-s3/main.tf'),
        ('Deploy application locally using Minikube', 'Done', 'deploy.sh, Makefile'),
        ('Store Docker image on DockerHub', 'Done', 'deepak415/csv-processor:latest'),
        ('Store code on GitHub', 'Done', 'github.com/AIINDEVOPS/silk'),
        ('Documentation and architecture diagram', 'Done', 'This document, README.md, ARCHITECTURE.md'),
    ]
)
page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — ARCHITECTURE DIAGRAM
# ════════════════════════════════════════════════════════════════════════════
h1('2. System Architecture Diagram')
para(
    'The diagram shows the complete system: kops cluster configuration (Task 1), '
    'and the Minikube local deployment with Nginx+Flask sidecar pod, MinIO S3 storage, '
    'HPA autoscaling, Helm packaging, and Ansible config management (Task 2).'
)
doc.add_paragraph()
img(os.path.join(SS, 'architecture-diagram.png'), width=Inches(6.2),
    caption_text='Figure 1 — Full system architecture diagram')

doc.add_paragraph()
h2('Traffic Flow (Minikube)')
bullet('User opens http://localhost:8080 (minikube tunnel maps LoadBalancer → 127.0.0.1)')
bullet('Nginx :80 receives request → serves /static/ from emptyDir, proxies / → Flask :5000')
bullet('Flask parses uploaded CSV → uploads to MinIO via boto3 → stores metadata.json')
bullet('Result page renders all rows in browser; homepage lists previously processed files')
bullet('HPA monitors CPU and Memory → scales pods 2→5 under load')

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — KOPS CLUSTER (TASK 1)
# ════════════════════════════════════════════════════════════════════════════
h1('3. Task 1 — Kubernetes Cluster Configuration (kops)', color=(0x8B,0x45,0x13))
para(
    'This task requires creating the kops cluster configuration files — a running '
    'cluster on AWS is not expected for this assignment. The configuration covers '
    'multiple instance groups with mixed lifecycle (spot and on-demand) and a '
    'Cluster Autoscaler for all worker instance groups.'
)

h2('3.1 Cluster Specification')
bullet('Cloud: AWS — us-east-1')
bullet('VPC CIDR: 172.20.0.0/16  |  Pod CIDR: 100.96.0.0/11  |  Service CIDR: 100.64.0.0/13')
bullet('Multi-AZ: us-east-1a, us-east-1b, us-east-1c')
bullet('CNI: Calico (supports NetworkPolicy)')
bullet('DNS: CoreDNS')

h2('3.2 Instance Groups — Multiple IGs, Mixed Lifecycle')
blue_header_table(
    ['Instance Group', 'Role', 'Instance Types', 'Min', 'Max', 'Lifecycle'],
    [
        ('master-us-east-1a', 'Control Plane', 'm5.large', '1', '1', 'On-Demand (fixed)'),
        ('master-us-east-1b', 'Control Plane', 'm5.large', '1', '1', 'On-Demand (fixed)'),
        ('master-us-east-1c', 'Control Plane', 'm5.large', '1', '1', 'On-Demand (fixed)'),
        ('nodes-ondemand', 'Workers', 'm5/m5a/m5n/m4.xlarge', '2', '10', 'On-Demand'),
        ('nodes-spot', 'Workers', 'm5/m4/r5/c5.xlarge', '0', '20', 'Spot (max $0.10/hr)'),
        ('nodes-gpu-spot', 'Workers', 'p3.2xl, p2.xl, g4dn.xl', '0', '5', 'Spot'),
    ],
    col_color='8B4513'
)

doc.add_paragraph()
h2('3.3 Cluster Autoscaler — All Instance Groups')
para(
    'The Cluster Autoscaler is configured for all worker instance groups. '
    'Masters are deliberately excluded — scaling masters would break etcd quorum '
    'and corrupt the cluster state.'
)
bullet('File: k8s-kops/cluster-autoscaler.yaml')
bullet('Auto-discovers worker IGs via ASG tags: k8s.io/cluster-autoscaler/enabled')
bullet('Expander: least-waste — picks the IG that wastes fewest CPU/memory resources')
bullet('Scale-down: removes idle nodes after 10 minutes')
bullet('Masters: NO autoscaler tags — fixed at 1 per AZ (etcd quorum protection)')

h2('3.4 Spot Instance Strategy')
bullet('Spot IG has PreferNoSchedule taint — used only when on-demand IG is full')
bullet('Multiple instance families in each spot IG for better spot pool availability')
bullet('GPU Spot IG has NoSchedule taint — requires explicit toleration on GPU workloads')

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — APPLICATION ON MINIKUBE (TASK 2)
# ════════════════════════════════════════════════════════════════════════════
h1('4. Task 2 — Application Deployment on Minikube', color=(0x2E,0x7D,0x32))
para(
    'The full application stack runs on Minikube locally. This covers all '
    'infrastructure requirements: Nginx+Flask pod, Service, HPA, Ansible, '
    'Helm, CSV web app, and S3 Glacier storage.'
)

# 4.1
h2('4.1 Nginx + Flask in the Same Pod (emptyDir, not NFS)')
para(
    'The assignment requires Nginx and the web application to run in the same pod, '
    'sharing public files (CSS, JS) through shared storage — explicitly not NFS.'
)
para('Solution: init container pattern with emptyDir volume.')
code_block(
'''Pod
├── Init Container (static-files-init)
│   └── Copies /app/static/* → emptyDir "shared-static"  (runs ONCE at pod start)
│
├── Container 1: nginx:1.25-alpine  [:80]
│   ├── proxy_pass http://127.0.0.1:5000  → Flask
│   └── Serves /static/* from emptyDir (not NFS — in-memory, same node)
│
└── Container 2: deepak415/csv-processor:latest  [:5000]
    ├── POST /upload  → parse CSV → upload MinIO → store metadata.json
    └── GET  /        → read metadata.json → list previously processed files

Shared Volumes (emptyDir — NOT NFS):
  shared-static:      CSS/JS files  (init → nginx)
  uploads-storage:    raw CSV uploads
  processed-storage:  metadata.json'''
)
para(
    'The emptyDir volume is allocated on the node where the pod runs — it is '
    'fast, in-memory where possible, and has zero external dependencies. '
    'This satisfies the "not nfs" requirement explicitly.'
)

# 4.2
h2('4.2 Service Object')
para('The application is exposed via a Kubernetes Service object of type LoadBalancer.')
code_block(
'''apiVersion: v1
kind: Service
metadata:
  name: csv-app-service
  namespace: csv-app
spec:
  type: LoadBalancer
  selector:
    app: csv-app
  ports:
    - port: 8080
      targetPort: 80    # Nginx inside the pod''')
para('On Minikube: run minikube tunnel → EXTERNAL-IP becomes 127.0.0.1:8080')
para('On AWS kops: AWS provisions a Network Load Balancer automatically')

# 4.3
h2('4.3 HPA Autoscaling')
para('Horizontal Pod Autoscaler scales the deployment between 2 and 5 replicas '
     'based on CPU and Memory utilisation.')
code_block(
'''minReplicas: 2
maxReplicas: 5
metrics:
  - CPU:    target averageUtilization: 70   # scale up above 70% CPU
  - Memory: target averageUtilization: 85   # scale up above 85% Memory
behavior:
  scaleUp:   stabilizationWindowSeconds: 60   # react fast to load
  scaleDown: stabilizationWindowSeconds: 300  # avoid flapping''')
para('Live state at idle: cpu:2%/70%  memory:59%/85% — both below thresholds → 2 replicas.')
para('Note: memory utilisation = actual usage / memory REQUEST (not limit). '
     'Request is set to 256Mi; actual idle ~195Mi = 76% → safely below 85% threshold.')

# 4.4
h2('4.4 Ansible Configuration Management')
para('Ansible manages application configuration using the kubernetes.core collection — '
     'no SSH into nodes, communicates with the cluster through the K8s API directly.')
code_block(
'''# ansible/site.yaml
- name: Manage Kubernetes application configuration
  hosts: localhost
  connection: local
  tasks:
    - name: Apply app ConfigMap
      kubernetes.core.k8s:
        state: present
        definition:
          kind: ConfigMap
          data:
            APP_NAME: "{{ app_name }}"
            STORAGE_BACKEND: "{{ storage_backend }}"
            AWS_REGION: "{{ aws_region }}"
            S3_BUCKET: "{{ s3_bucket }}"

    - name: Apply app Secret
      kubernetes.core.k8s:
        definition:
          kind: Secret
          stringData:
            S3_BUCKET: "{{ s3_bucket }}"
            SECRET_KEY: "{{ vault_secret_key }}"

    - name: Wait for rollout
      kubernetes.core.k8s_rollout_status:
        kind: Deployment
        name: csv-app

    - name: Check Flask health endpoint
      kubernetes.core.k8s_exec:
        command: ["curl", "-s", "http://localhost:5000/health"]''')
code_block(
'''# Install and run
ansible-galaxy collection install -r ansible/requirements.yml
ansible-playbook ansible/site.yaml -i ansible/inventory/k8s.yaml''')

# 4.5
h2('4.5 Helm Multi-Environment Packaging')
para('A single Helm chart is used to deploy the application to any environment. '
     'Environment-specific settings are in separate values files.')
code_block(
'''helm/
├── csv-app/                  Single reusable chart
│   ├── Chart.yaml
│   ├── values.yaml           Defaults
│   └── templates/
│       ├── deployment.yaml   replicas omitted when HPA enabled
│       ├── service.yaml
│       ├── hpa.yaml
│       ├── configmap.yaml
│       ├── pdb.yaml
│       └── namespace.yaml
└── environments/
    ├── local-values.yaml     Minikube: LoadBalancer :8080, MinIO, 256Mi
    ├── dev-values.yaml       Dev cluster: NodePort, AWS S3, 1 replica
    └── prod-values.yaml      Prod: LoadBalancer :80, AWS S3, 4 replicas''')
code_block(
'''# Deploy to local (Minikube)
helm upgrade --install csv-app helm/csv-app \\
  -f helm/environments/local-values.yaml \\
  --namespace csv-app --create-namespace --wait

# Deploy to prod (kops cluster)
helm upgrade --install csv-app helm/csv-app \\
  -f helm/environments/prod-values.yaml \\
  --namespace csv-app --create-namespace --wait''')

# 4.6
h2('4.6 CSV Web Application')
para('A Python Flask application that parses CSV files and displays results in the browser.')
numbered('User opens http://localhost:8080 — sees upload form and previously processed files')
numbered('User drags CSV file onto the drop zone or clicks Browse File')
numbered('Flask receives the file, prefixes with timestamp, parses with csv.reader')
numbered('All rows uploaded to MinIO at processed/YYYY/MM/DD/<timestamp>_<filename>.csv')
numbered('Result page renders all 751 rows (Product ID, Name, Price) as an HTML table')
numbered('Metadata (filename, rows, path, timestamp) stored in metadata.json')
numbered('Homepage GET / reads metadata.json and lists all previously processed files')

doc.add_paragraph()
blue_header_table(
    ['Endpoint', 'Method', 'Description'],
    [
        ('/', 'GET', 'Homepage: upload form + previously processed files list'),
        ('/upload', 'POST', 'Receive CSV, parse, upload to MinIO, render result table'),
        ('/health', 'GET', 'Health check: {"status":"healthy","storage_backend":"minio"}'),
        ('/ready', 'GET', 'Readiness probe for Kubernetes'),
    ]
)

# 4.7
h2('4.7 S3 Glacier Transition (Terraform)')
para(
    'The assignment explicitly requires: "Waiting you to implement s3 glacier '
    'transition on s3 config." This is implemented in terraform-s3/main.tf.'
)
para('Lifecycle rule: csv-glacier-transition — applies to all objects under processed/ prefix.')
doc.add_paragraph()

blue_header_table(
    ['Day', 'Storage Class', 'Cost vs Standard', 'Retrieval Time'],
    [
        ('Day 0',    'STANDARD',     '100%',  'Instant'),
        ('Day 30',   'STANDARD_IA',  '~60%',  'Instant (per-retrieval fee)'),
        ('Day 90',   'GLACIER_IR',   '~32%',  'Milliseconds'),
        ('Day 180',  'GLACIER',      '~20%',  '3–5 hours'),
        ('Day 365',  'DEEP_ARCHIVE', '~5%',   '12 hours'),
        ('Day 2555', 'DELETE',       '—',     'Compliance window (7 years)'),
    ],
    col_color='FF9900'
)
doc.add_paragraph()
code_block(
'''# terraform-s3/main.tf — glacier transition rule
resource "aws_s3_bucket_lifecycle_configuration" "csv_uploads" {
  rule {
    id     = "csv-glacier-transition"
    status = "Enabled"
    filter { prefix = "processed/" }

    transition { days = 30  storage_class = "STANDARD_IA"  }
    transition { days = 90  storage_class = "GLACIER_IR"   }
    transition { days = 180 storage_class = "GLACIER"      }
    transition { days = 365 storage_class = "DEEP_ARCHIVE" }
    expiration  { days = 2555 }                # delete after 7 years
  }
}''')
code_block(
'''cd terraform-s3
terraform init
terraform plan
terraform apply''')

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — STEP BY STEP LOCAL DEPLOY
# ════════════════════════════════════════════════════════════════════════════
h1('5. Local Deploy — Step by Step')
para('Full step-by-step deployment on macOS using Minikube and the docker driver.')

h2('Prerequisites')
code_block(
'''# macOS — install tools
brew install minikube kubectl helm ansible
brew install --cask docker    # Docker Desktop

# Start Docker Desktop (required before any Minikube commands)''')

h2('Step 1 — Start Minikube')
code_block(
'''minikube start --driver=docker --cpus=4 --memory=8192
minikube addons enable metrics-server    # required for HPA

# Verify
kubectl get nodes
# NAME       STATUS   ROLES           AGE   VERSION
# minikube   Ready    control-plane   30s   v1.29.x''')

h2('Step 2 — Build Docker image inside Minikube')
code_block(
'''# Build directly into Minikube — no pull from DockerHub needed
eval $(minikube docker-env)
docker build -t csv-processor:local ./app
eval $(minikube docker-env -u)   # reset to host docker''')

h2('Step 3 — Deploy MinIO (local S3)')
code_block(
'''kubectl apply -f local/k8s/namespace.yaml
kubectl apply -f local/k8s/minio.yaml
kubectl rollout status deployment/minio -n csv-app --timeout=120s

# Create the csv-uploads bucket
kubectl apply -f local/k8s/minio-init-job.yaml
kubectl wait --for=condition=complete job/minio-setup -n csv-app --timeout=120s''')

h2('Step 4 — Deploy application via Helm')
code_block(
'''helm upgrade --install csv-app helm/csv-app \\
  -f helm/environments/local-values.yaml \\
  --namespace csv-app --create-namespace --wait --timeout 5m

# Verify — should show 2/2 Running
kubectl get pods -n csv-app
# NAME                 READY   STATUS    RESTARTS   AGE
# csv-app-xxx-yyy      2/2     Running   0          60s   <- Nginx + Flask
# minio-xxx-aaa        1/1     Running   0          90s''')

h2('Step 5 — Start minikube tunnel (separate terminal)')
code_block(
'''# Open a NEW terminal — keep it open
minikube tunnel
# csv-app-service EXTERNAL-IP will become 127.0.0.1''')

h2('Step 6 — Open application')
code_block(
'''open http://localhost:8080

# MinIO console (S3 browser)
minikube service minio-console -n csv-app
# Login: minioadmin / minioadmin''')

h2('Step 7 — Upload CSV and verify')
numbered('Open http://localhost:8080')
numbered('Drag tasks/soh-1-.csv onto the upload zone (751 fashion products)')
numbered('Click Upload & Process')
numbered('Result page shows 751 rows: Product ID, Name, Price')
numbered('MinIO console → csv-uploads → processed/2026/06/06/ shows the uploaded file')
numbered('Return to homepage — previously processed files list updated')

h2('Step 8 — Run Ansible')
code_block(
'''ansible-galaxy collection install -r ansible/requirements.yml
export KUBECONFIG=~/.kube/config
ansible-playbook ansible/site.yaml -i ansible/inventory/k8s.yaml -v''')

h2('Step 9 — Verify all resources')
code_block(
'''kubectl get pods,svc,hpa -n csv-app -o wide

# Expected:
# pod/csv-app-xxx   2/2  Running                              <- 2 containers
# csv-app-service   LoadBalancer  127.0.0.1  8080:xxxxx/TCP   <- LoadBalancer
# csv-app-hpa       cpu:2%/70%  memory:59%/85%  min:2 max:5  <- HPA healthy''')

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — DEPLOY.SH
# ════════════════════════════════════════════════════════════════════════════
h1('6. One-Command Deploy (deploy.sh)')
para(
    'A single shell script automates the entire deployment. Anyone can clone the '
    'repository and run one command to get the full stack running locally.'
)
code_block(
'''# Clone and deploy
git clone https://github.com/AIINDEVOPS/silk.git
cd silk
bash deploy.sh

# The script will:
#   1. Check Docker Desktop is running
#   2. Install minikube/kubectl/helm if missing (macOS via brew)
#   3. Start Minikube with metrics-server addon
#   4. Build Docker image inside Minikube (no pull needed)
#   5. Deploy MinIO and create the csv-uploads bucket
#   6. Deploy app via Helm (local-values.yaml)
#   7. Print access instructions

# After script completes:
minikube tunnel          # run in separate terminal
open http://localhost:8080

# Tear down
bash deploy.sh --clean''')

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — SCREENSHOTS
# ════════════════════════════════════════════════════════════════════════════
h1('7. Application Screenshots')
para('All screenshots captured from the live Minikube deployment at http://localhost:8080.')

h2('7.1 Upload Form — Application Homepage')
para(
    'CSV File Processor running at http://localhost:8080 via minikube tunnel. '
    'The banner shows "Storage: MinIO (local Minikube)" confirming the local S3 backend. '
    'The drag-drop upload zone satisfies the assignment requirement: '
    '"basic interface to upload CSV and show previously processed files."'
)
img(os.path.join(SS, 'CSV_File_Processor_1.png'),
    caption_text='Figure 2 — Application homepage with upload form at http://localhost:8080')

doc.add_paragraph()

h2('7.2 File Selected — Ready to Process')
para(
    'soh-1-.csv (the assignment-provided CSV with 751 fashion products) selected '
    'via the drag-drop interface, ready to click Upload & Process.'
)
img(os.path.join(SS, 'CSV_File_Processor_2.png'),
    caption_text='Figure 3 — soh-1-.csv (751 rows) selected and ready to upload')

page_break()

h2('7.3 CSV Processed — 751 Rows Displayed')
para(
    'Result page after processing. Shows: filename with timestamp, total rows (751), '
    'storage backend (MinIO local), and the full MinIO storage path. '
    'The product table displays all 751 fashion items with Product ID, Name, and Price. '
    'This satisfies: "parse and process CSV files — print content of the lines to browser."'
)
img(os.path.join(SS, 'CSV_File_Processor_3.png'),
    caption_text='Figure 4 — 751 rows parsed and displayed; file stored at MinIO processed/ path')

doc.add_paragraph()

h2('7.4 Kubernetes Resources — Pods, Service, HPA')
para(
    'Terminal output of kubectl get pods,svc,hpa -n csv-app showing: '
    '2/2 Running pods (Nginx + Flask sidecar containers in same pod), '
    'LoadBalancer service with EXTERNAL-IP 127.0.0.1:8080 via minikube tunnel, '
    'and HPA with cpu:2%/70% and memory:59%/85% at idle (below both thresholds).'
)
img(os.path.join(SS, 'pods-svc-hpa-status-kubectl.png'),
    caption_text='Figure 5 — kubectl output: 2/2 pods running, LoadBalancer active, HPA healthy')

page_break()

h2('7.5 MinIO Console — Login')
para(
    'MinIO Object Store web console — the local S3-compatible storage running in Minikube. '
    'This serves as the local equivalent of AWS S3 for development. '
    'Credentials: minioadmin / minioadmin.'
)
img(os.path.join(SS, 'Minio_Storage_1.png'), width=Inches(4.5),
    caption_text='Figure 6 — MinIO Object Store login (local S3-compatible storage in Minikube)')

doc.add_paragraph()

h2('7.6 MinIO Bucket — csv-uploads')
para(
    'The csv-uploads bucket with the processed/ prefix. All CSV files are stored '
    'under processed/YYYY/MM/DD/ — the same path structure targeted by the '
    'Terraform S3 Glacier lifecycle rules in production. '
    '223.3 KiB total, 6 objects.'
)
img(os.path.join(SS, 'Minio_Storage_2.png'),
    caption_text='Figure 7 — MinIO csv-uploads bucket with processed/ prefix (223.3 KiB, 6 objects)')

page_break()

h2('7.7 MinIO — Processed CSV Files by Date')
para(
    'Contents of processed/2026/06/06/ showing four uploaded CSV files with '
    'timestamp-prefixed names. This path structure matches the boto3 upload '
    'path in app.py and mirrors the Glacier lifecycle prefix in Terraform.'
)
img(os.path.join(SS, 'Minio_Storage_3.png'),
    caption_text='Figure 8 — Processed CSVs at processed/2026/06/06/ — date-partitioned storage')

doc.add_paragraph()

h2('7.8 Minikube Dashboard — Workload Status')
para(
    'Kubernetes dashboard filtered to csv-app namespace showing: '
    '2 Deployments running (csv-app + minio), 3 Pods running, '
    '4 Replica Sets (current + rollout history), 1 completed Job (MinIO init).'
)
img(os.path.join(SS, 'minikube-dashboard-status.png'),
    caption_text='Figure 9 — Minikube dashboard: csv-app namespace workload status')

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 — REPO STRUCTURE
# ════════════════════════════════════════════════════════════════════════════
h1('8. Repository Structure')
code_block(
'''.
├── deploy.sh                     One-command deploy (clone → full stack)
├── Makefile                      Individual commands for each step
│
├── app/                          Web application
│   ├── app.py                    Flask: CSV parse, MinIO/S3 upload, metadata
│   ├── Dockerfile                python:3.12-slim, non-root user uid 1001
│   ├── nginx.conf                Reverse proxy + emptyDir static serving
│   ├── requirements.txt          Flask, boto3, gunicorn, werkzeug
│   ├── templates/
│   │   ├── index.html            Upload form + previously processed files list
│   │   └── result.html           Full CSV table (751 rows)
│   └── static/css/ + js/         Served via emptyDir (not NFS)
│
├── k8s-kops/                     TASK 1 — kops cluster configuration
│   ├── cluster.yaml              Cluster spec: VPC, subnets, Calico
│   ├── instancegroups.yaml       3 masters (fixed) + 3 worker IGs
│   ├── cluster-autoscaler.yaml   CA for worker IGs (least-waste expander)
│   ├── deployment.yaml           App deployment (nginx+flask sidecar, emptyDir)
│   └── service-hpa.yaml          LoadBalancer + HPA + PodDisruptionBudget
│
├── helm/                         TASK 2 — Helm packaging
│   ├── csv-app/                  Reusable chart (deployment, service, hpa, cm, pdb)
│   └── environments/
│       ├── local-values.yaml     Minikube: LoadBalancer :8080, MinIO, 256Mi req
│       ├── dev-values.yaml       Dev: NodePort, AWS S3, 1 replica
│       └── prod-values.yaml      Prod: LoadBalancer, AWS S3, 4 replicas
│
├── ansible/                      TASK 2 — Config management
│   ├── site.yaml                 ConfigMap + Secret + Deployment patch + health
│   ├── requirements.yml          kubernetes.core >= 3.0.0
│   ├── inventory/k8s.yaml        localhost, ansible_connection: local
│   └── group_vars/all.yaml       App vars, AWS region, image reference
│
├── terraform-s3/                 TASK 2 — S3 + Glacier lifecycle
│   ├── main.tf                   Bucket + glacier transitions + IAM (local state)
│   ├── variables.tf
│   └── outputs.tf
│
├── local/                        Local Minikube manifests and helpers
│   ├── k8s/                      namespace, minio, minio-init-job, deployment
│   ├── ansible/                  Local Ansible playbook + inventory
│   └── scripts/smoke-test.sh     End-to-end smoke test
│
├── screenshots/                  All screenshots for documentation
├── docker-compose.yml            Quick start without Minikube
├── README.md                     Project overview and quick-start
├── ARCHITECTURE.md               Full architecture + Mermaid diagram
└── LOCAL-TESTING-GUIDE.md        Detailed Minikube walkthrough''')

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 9 — TECH STACK
# ════════════════════════════════════════════════════════════════════════════
h1('9. Technology Stack')
blue_header_table(
    ['Layer', 'Technology', 'Purpose'],
    [
        ('Web Framework',     'Python 3.12 + Flask',           'CSV parsing and HTML rendering'),
        ('Web Server',        'Nginx 1.25-alpine',             'Reverse proxy + static files (emptyDir)'),
        ('WSGI Server',       'Gunicorn',                      'Production multi-worker server'),
        ('Container',         'Docker 25+',                    'Immutable image — deepak415/csv-processor'),
        ('Cluster Config',    'kops 1.29+',                    'Declarative K8s cluster definition on AWS'),
        ('Local Cluster',     'Minikube (docker driver)',       'Local deployment on Mac'),
        ('Node Autoscaling',  'Cluster Autoscaler',            'Scales worker IGs based on pending pods'),
        ('Pod Autoscaling',   'HPA v2',                        'Scales pods on CPU + Memory metrics'),
        ('Helm',              'Helm 3.14+',                    'Multi-environment K8s packaging'),
        ('Config Management', 'Ansible + kubernetes.core 3+',  'ConfigMap/Secret management via K8s API'),
        ('IaC',               'Terraform 1.7+',                'S3 bucket + Glacier lifecycle + IAM'),
        ('Local Storage',     'MinIO',                         'S3-compatible local object store'),
        ('Cloud Storage',     'AWS S3 + Glacier',              'Production CSV archiving'),
        ('Image Registry',    'DockerHub',                     'deepak415/csv-processor:latest'),
        ('Code Storage',      'GitHub',                        'github.com/AIINDEVOPS/silk'),
        ('CNI',               'Calico',                        'NetworkPolicy support on kops'),
    ]
)

doc.add_paragraph()
doc.add_paragraph()
divider()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(
    'DevOps Case Study  |  deepak415/csv-processor  |  '
    'github.com/AIINDEVOPS/silk  |  June 2026'
)
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0x6C,0x75,0x7D)
fr.italic = True

doc.save(OUT)
print(f"Saved: {OUT}")
